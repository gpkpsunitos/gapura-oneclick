#!/usr/bin/env python3

from __future__ import annotations

import argparse
import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def compact(text: str) -> str:
    return " ".join(text.split())


def xml_counts(archive: zipfile.ZipFile) -> dict[str, int]:
    document_xml = etree.fromstring(archive.read("word/document.xml"))
    counts = {
        "insertions": len(document_xml.xpath(".//w:ins", namespaces=NS)),
        "deletions": len(document_xml.xpath(".//w:del", namespaces=NS)),
        "field_chars": len(document_xml.xpath(".//w:fldChar", namespaces=NS)),
        "field_instructions": len(document_xml.xpath(".//w:instrText", namespaces=NS)),
        "comment_range_starts": len(document_xml.xpath(".//w:commentRangeStart", namespaces=NS)),
        "comment_range_ends": len(document_xml.xpath(".//w:commentRangeEnd", namespaces=NS)),
        "comment_references": len(document_xml.xpath(".//w:commentReference", namespaces=NS)),
    }
    if "word/comments.xml" in archive.namelist():
        comments_xml = etree.fromstring(archive.read("word/comments.xml"))
        counts["comments"] = len(comments_xml.xpath(".//w:comment", namespaces=NS))
    else:
        counts["comments"] = 0
    return counts


def section_record(index, section):
    return {
        "index": index,
        "page_width_emu": section.page_width,
        "page_height_emu": section.page_height,
        "top_margin_emu": section.top_margin,
        "bottom_margin_emu": section.bottom_margin,
        "left_margin_emu": section.left_margin,
        "right_margin_emu": section.right_margin,
        "header_distance_emu": section.header_distance,
        "footer_distance_emu": section.footer_distance,
        "start_type": str(section.start_type),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document(args.input)
    with zipfile.ZipFile(args.input) as archive:
        package = {
            "parts": sorted(archive.namelist()),
            "media": sorted(name for name in archive.namelist() if name.startswith("word/media/")),
            "xml_counts": xml_counts(archive),
        }

    paragraphs = []
    captions = []
    headings = []
    for index, paragraph in enumerate(document.paragraphs):
        text = compact(paragraph.text)
        style = paragraph.style.name if paragraph.style else None
        record = {
            "index": index,
            "text": text,
            "style": style,
            "runs": len(paragraph.runs),
            "has_drawing": bool(paragraph._p.xpath(".//*[local-name()='drawing']")),
            "has_page_break": bool(paragraph._p.xpath(".//*[local-name()='br' and @*[local-name()='type']='page']")),
        }
        paragraphs.append(record)
        if style and style.startswith("Heading"):
            headings.append(record)
        if text.lower().startswith(("gambar ", "tabel ")):
            captions.append(record)

    tables = []
    for index, table in enumerate(document.tables):
        rows = []
        for row_index, row in enumerate(table.rows):
            rows.append({
                "index": row_index,
                "cells": [compact(cell.text) for cell in row.cells],
            })
        tables.append({
            "index": index,
            "style": table.style.name if table.style else None,
            "row_count": len(table.rows),
            "column_count": len(table.columns),
            "rows": rows,
        })

    inventory = {
        "source": str(args.input.resolve()),
        "sha256": sha256(args.input),
        "size_bytes": args.input.stat().st_size,
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "inline_shape_count": len(document.inline_shapes),
        "paragraphs": paragraphs,
        "headings": headings,
        "captions": captions,
        "tables": tables,
        "sections": [section_record(index, section) for index, section in enumerate(document.sections)],
        "package": package,
    }

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(inventory, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({
        "sha256": inventory["sha256"],
        "paragraphs": inventory["paragraph_count"],
        "tables": inventory["table_count"],
        "sections": inventory["section_count"],
        "inline_shapes": inventory["inline_shape_count"],
        "captions": len(captions),
        **package["xml_counts"],
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Apply implementation-aligned revisions and true Word comments to the thesis."""

from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


AUTHOR = "Codex — Audit Implementasi"
INITIALS = "CAI"
FIVE_DIVISIONS = "DIVISI_OS, DIVISI_OP, DIVISI_HT, DIVISI_UQ, dan DIVISI_OT"


PARAGRAPH_REPLACEMENTS = {
155: "Pelaporan irregularity pada operasional ground handling membutuhkan alur data yang cepat, konsisten, dan mudah ditelusuri karena setiap penyimpangan layanan dapat memengaruhi kualitas pelayanan penerbangan. Pada sistem berjalan, proses pencatatan masih tersebar pada Microsoft Forms, berkas Microsoft Excel, Google Sheets, dan dasbor Looker Studio. Fragmentasi tersebut menimbulkan keterlambatan sinkronisasi, meningkatkan risiko kesalahan pencatatan, dan membatasi pemanfaatan data historis sebagai dasar analisis operasional. Penelitian ini merancang dan mengimplementasikan sistem pelaporan irregularity ground handling terintegrasi berbasis arsitektur mikrolayanan. Pada implementasi, laporan baru ditulis terlebih dahulu ke Google Sheets, lalu identitas dan metadatanya disinkronkan secara best-effort ke Supabase PostgreSQL untuk pembacaan aplikasi, komentar, sesi, dan audit. Bukti laporan disimpan di Google Drive, sedangkan dokumen akhir DOCX/PDF menggunakan alur Supabase Storage yang terpisah. Metode pengembangan mencakup analisis kebutuhan, perancangan arsitektur, pembangunan antarmuka menggunakan Next.js, integrasi layanan FastAPI, serta pengujian fungsional dan evaluasi pada lingkungan lokal. Kapabilitas analitik yang diuji meliputi peramalan volume laporan, analisis musiman, klasifikasi subkategori, identifikasi akar masalah, dan penghitungan skor risiko. Evaluasi terhadap 1.146 rekaman aktif menunjukkan bahwa fungsi utama sistem dapat dijalankan pada lingkungan uji. Model peramalan memperoleh MAE 0,974 dan RMSE 1,2242, sedangkan akurasi validasi silang untuk klasifikasi kategori, subkategori, dan akar masalah berturut-turut sebesar 0,6803, 0,5729, dan 0,4381. Hasil tersebut menunjukkan kelayakan fungsional purwarupa; keluaran analitik tetap memerlukan verifikasi pengguna berwenang sebelum digunakan dalam evaluasi operasional.",
160: "Irregularity reporting in ground handling operations requires a fast, consistent, and traceable data flow because every service deviation can affect flight-service quality. In the existing process, records remain distributed across Microsoft Forms, Microsoft Excel files, Google Sheets, and a Looker Studio dashboard. This fragmentation delays synchronization, increases recording-error risk, and limits the use of historical data for operational analysis. This study designs and implements an integrated microservice-based ground-handling irregularity reporting system. In the implemented flow, each new report is written to Google Sheets first; its identity and metadata are then synchronized on a best-effort basis to Supabase PostgreSQL for application reads, comments, sessions, and audit functions. Report evidence is stored in Google Drive, while final DOCX/PDF documents use a separate Supabase Storage flow. The development method covers requirements analysis, architecture design, a Next.js interface, FastAPI integration, functional testing, and local-environment evaluation. The evaluated analytics include report-volume forecasting, seasonality analysis, subcategory classification, root-cause identification, and risk scoring. Evaluation of 1,146 active records shows that the main functions run in the test environment. The forecast obtained an MAE of 0.974 and an RMSE of 1.2242, while cross-validation accuracy for category, subcategory, and root-cause classification was 0.6803, 0.5729, and 0.4381. These results support the prototype's functional feasibility; authorized users must still verify analytic outputs before operational use.",
269: "Untuk menangani permasalahan tersebut, penelitian ini merancang sistem pelaporan irregularity terintegrasi berbasis web. Implementasi mempertahankan Google Sheets sebagai tujuan penulisan pertama laporan dan sumber operasional layanan analitik. Identitas serta metadata laporan kemudian disinkronkan secara best-effort ke Supabase PostgreSQL untuk kebutuhan pembacaan aplikasi, komentar, sesi, audit, dan fitur terkait. Arsitektur mikrolayanan memisahkan tanggung jawab layanan serta mendukung modularitas dan skalabilitas. Sistem juga menerapkan Natural Language Processing (NLP) untuk klasifikasi akar masalah dan peramalan untuk menaksir volume laporan. Kedua kapabilitas tersebut diposisikan sebagai dukungan analitik bagi evaluasi operasional.",
278: "Solusi yang diajukan adalah aplikasi pelaporan irregularity berbasis web dengan alur write-first ke Google Sheets dan sinkronisasi metadata ke Supabase PostgreSQL. Arsitektur mikrolayanan memisahkan modul transaksi pelaporan dari modul analitik yang mencakup klasifikasi otomatis dan peramalan volume laporan. Pemisahan tersebut memungkinkan alur pelaporan yang terintegrasi sekaligus menghasilkan informasi pendukung evaluasi operasional.",
281: "Mengembangkan sistem pelaporan irregularity berbasis web yang menulis laporan ke Google Sheets terlebih dahulu dan menyinkronkan identitas serta metadata ke Supabase PostgreSQL untuk mengurangi fragmentasi data.",
307: "Penelitian ini menerapkan integrasi data dengan membedakan sumber operasional dan read model aplikasi. Laporan baru ditulis terlebih dahulu ke Google Sheets karena alur tersebut masih menjadi integrasi transaksi yang aktif. Setelah identitas laporan diperoleh, metadata disinkronkan secara best-effort ke tabel ground_handling_irregularity_report pada Supabase PostgreSQL untuk pembacaan aplikasi, filter, komentar, dashboard, dan metadata terkait. Dengan demikian, Supabase tidak diposisikan sebagai tujuan penulisan pertama laporan pada implementasi yang diaudit.",
313: "Basis data adalah kumpulan data yang tersusun logis untuk memudahkan penyimpanan, pencarian, dan pengelolaan. Pada penelitian ini, Supabase PostgreSQL menyimpan akun aplikasi, sesi, komentar, metadata bukti, dokumen, audit, serta read model laporan yang disinkronkan. Google Sheets tetap menjadi tujuan penulisan pertama laporan dan sumber operasional yang dibaca layanan analitik.",
317: "Autentikasi memastikan identitas pengguna, sedangkan otorisasi menentukan hak akses terhadap fungsi dan data. Aplikasi memverifikasi kata sandi bcrypt pada data pengguna yang dikelola aplikasi, kemudian menerbitkan cookie JWT bertanda tangan. Status sesi, masa berlaku, aktivitas terakhir, dan pencabutan juga dicatat pada tabel security_sessions sehingga implementasi bukan sesi stateless murni. Pemeriksaan peran dilakukan pada halaman dan API sesuai kewenangan yang diimplementasikan.",
318: "Jejak audit melengkapi autentikasi dan otorisasi dengan mencatat aktivitas penting yang benar-benar diinstrumentasikan. Pada skema live, audit aplikasi disimpan pada audit_logs, sedangkan cache analitik disimpan pada ai_cache_entries; tabel ai_audit_logs yang disebut pada rancangan lama tidak terdapat pada skema live yang diaudit.",
422: f"Kebutuhan fungsional disusun berdasarkan aktor yang berinteraksi langsung dengan sistem. Aktor penelitian mencakup STAFF_CABANG, MANAGER_CABANG, {FIVE_DIVISIONS}, DIVISI_OCS, ANALYST, SUPER_ADMIN, pelapor publik pada alur pengajuan publik, dan sistem itu sendiri. {FIVE_DIVISIONS} tetap dimodelkan sebagai lima aktor terpisah dengan fitur, izin, dan alur kerja yang sama. DIVISI_OCS memiliki workspace tersendiri.",
439: FIVE_DIVISIONS,
443: f"Masing-masing dari {FIVE_DIVISIONS} memantau laporan, membuka detail, membaca dashboard, memberi komentar atau tindak lanjut, serta menggunakan analisis AI dan asisten virtual dengan pola akses yang sama.",
445: "DIVISI_OCS",
446: "Melakukan login ke aplikasi dan membuka workspace OCS yang terpisah.",
447: "Melihat serta mengelola data tab dan tautan cepat pada workspace OCS sesuai hak akses yang diimplementasikan.",
448: "Mengakses fungsi pendamping yang tersedia bagi OCS, termasuk asisten virtual, tanpa menyamakan workspace OCS dengan halaman lima peran divisi lainnya.",
449: "Mempertahankan pemisahan data dan navigasi workspace OCS dari halaman pemantauan laporan umum.",
467: "Mengelola status akun dan peran melalui fungsi administrasi yang tersedia; mutasi data master yang belum memiliki rute implementasi tidak dinyatakan sebagai fitur selesai.",
474: "Menulis laporan baru ke Google Sheets terlebih dahulu, kemudian menyinkronkan identitas dan metadata laporan ke Supabase PostgreSQL secara best-effort.",
475: "Mengunggah bukti ke Google Drive dan menyimpan ledger kepemilikan serta pengaitan bukti pada Supabase PostgreSQL; dokumen akhir DOCX/PDF menggunakan alur Supabase Storage yang terpisah.",
476: "Membaca Google Sheets sebagai sumber operasional layanan analitik dan menggunakan Supabase PostgreSQL sebagai read model aplikasi yang tersinkron.",
477: f"Menampilkan laporan kepada manajer cabang, {FIVE_DIVISIONS}, analis, atau super admin sesuai pemeriksaan akses yang diimplementasikan.",
493: "Rancangan basis data memetakan data aplikasi dan read model operasional. Metadata skema live Supabase menunjukkan tabel pengguna, stasiun, laporan tersinkron, komentar, bukti, sesi, audit, serta dokumen. Google Sheets tetap menjadi tujuan penulisan pertama laporan. ERD pada Gambar III.1 memakai notasi Chen; relasi tanpa foreign key fisik ditandai sebagai asosiasi logis aplikasi.",
498: "Gambar III.2 menunjukkan struktur logis tabel terpilih berdasarkan primary key, foreign key, nullability, dan asosiasi aplikasi pada skema live Supabase.",
503: "Gambar III.3 menunjukkan hubungan antarkomponen. Aplikasi menulis laporan terlebih dahulu ke Google Sheets, lalu menyinkronkan identitas dan metadata ke Supabase PostgreSQL secara best-effort. Bukti laporan diunggah ke Google Drive, sedangkan berkas akhir DOCX/PDF memakai Supabase Storage melalui alur report_documents yang terpisah. Supabase PostgreSQL juga menyimpan akun aplikasi, sesi, komentar, audit, dan read model laporan.",
507: "Aplikasi OneClick mengorkestrasi permintaan analitik ke layanan ML FastAPI dan membaca data operasional dari Google Sheets. Endpoint /api/ai/analyze menerima satu narasi, sedangkan /api/ai/analyze-all saat ini merupakan alias GET untuk ikhtisar agregat. Untuk asisten virtual, OneClick memvalidasi sesi dan menerapkan kuota lima permintaan per hari sebelum meneruskan pertanyaan ke Gapura RAG. Gapura RAG mengembalikan jawaban dengan sumber dan halaman; otorisasi dokumen per peran atau divisi belum diklaim karena metadata chunk belum memuat kontrol tersebut.",
509: f"Use case diagram pada Gambar III.4 memodelkan STAFF_CABANG, MANAGER_CABANG, {FIVE_DIVISIONS}, DIVISI_OCS, ANALYST, dan SUPER_ADMIN. Kelima peran divisi tetap ditampilkan terpisah dengan hubungan yang sama, sedangkan OCS memakai workspace tersendiri.",
512: f"STAFF_CABANG mengajukan laporan dan memantau laporan miliknya. MANAGER_CABANG memantau laporan dalam cakupan cabang dan menyetujui akun staf. {FIVE_DIVISIONS} memiliki fitur yang sama untuk pemantauan, tindak lanjut, analisis AI, dan asisten virtual. DIVISI_OCS mengakses workspace terpisah. ANALYST melakukan analisis lintas cabang, sedangkan SUPER_ADMIN mengelola pengguna, sesi, dan audit sesuai fungsi yang tersedia.",
518: "Gambar III.6 menunjukkan pengajuan laporan: pelapor mengisi formulir, bukti diunggah ke Google Drive, laporan ditulis ke Google Sheets, kemudian metadata disinkronkan ke Supabase PostgreSQL secara best-effort.",
527: "Gambar III.8 menunjukkan alur analisis AI dan asisten virtual. Sistem memvalidasi sesi dan masukan; khusus RAG, OneClick juga memeriksa kuota lima permintaan per hari. Hasil analitik atau jawaban beserta sumber ditampilkan untuk diverifikasi pengguna.",
538: "Gambar III.10 menunjukkan urutan pengajuan laporan: bukti diunggah ke Google Drive, data laporan ditulis ke Google Sheets, lalu metadata disinkronkan ke Supabase PostgreSQL tanpa membatalkan penulisan Sheets apabila sinkronisasi metadata gagal.",
548: "Gambar III.12 menunjukkan urutan analisis AI dan RAG. Setelah validasi sesi dan kuota yang relevan, OneClick memanggil layanan ML atau Gapura RAG, kemudian menampilkan hasil analitik atau jawaban beserta sumber dan halaman untuk verifikasi pengguna.",
565: f"Tahap pembuatan kode menerjemahkan kebutuhan STAFF_CABANG, MANAGER_CABANG, {FIVE_DIVISIONS}, DIVISI_OCS, ANALYST, SUPER_ADMIN, pelapor publik, dan sistem menjadi halaman, API, validasi, integrasi penyimpanan, serta layanan analitik. Kelima peran divisi memiliki pola fitur yang sama, sementara OCS memiliki workspace terpisah.",
573: "Sumber data layanan analitik berupa Google Sheets yang menjadi sumber operasional aktif untuk laporan. Layanan membaca tab NON CARGO dan CGO dengan akun layanan Google berizin baca-saja. Kredensial diinjeksi melalui variabel lingkungan. Snapshot yang dievaluasi berjumlah 1.149 baris gabungan sebelum penyaringan dan 1.146 rekaman aktif untuk evaluasi; Supabase menyimpan read model laporan yang tersinkron, bukan salinan asal yang menulis ke Sheets.",
589: "Data pelatihan model analitik dibaca dari Google Sheets sebagai sumber operasional. Layanan membentuk kerangka data dan skema kolom melalui data_fetcher dan schema_detector. Kolom masukan mencakup tanggal, kategori, maskapai, stasiun, area, formulir sumber, uraian laporan, akar masalah hasil investigasi, dan status. Untuk mencegah kebocoran target, kolom pembentuk label dikeluarkan dari fitur pada tugas terkait.",
610: "Asisten virtual diimplementasikan sebagai layanan mikro Gapura RAG berbasis FastAPI/Uvicorn. Aplikasi OneClick menjadi gerbang pengguna: sesi aplikasi diverifikasi dan kuota lima pertanyaan per hari diterapkan sebelum permintaan diteruskan. Pipeline RAG melakukan ekstraksi, chunking, embedding, retrieval Pinecone, reranking, dan pembangkitan jawaban yang menyertakan sumber serta halaman. Antarmuka langsung layanan tetap merupakan komponen teknis terpisah.",
618: "Ketahanan layanan RAG dibangun melalui pembatasan laju pada layanan, validasi payload dan berkas, serta kuota lima permintaan per akun per hari pada proxy OneClick. Jika LLM tidak tersedia, respons kesalahan dikendalikan dan metadata retrieval dapat tetap disertakan. Metadata chunk saat ini belum mengodekan ACL dokumen per peran atau divisi; kondisi tersebut dinyatakan sebagai keterbatasan, bukan sebagai kontrol yang sudah diterapkan.",
625: f"Gambar III.14 menunjukkan contoh halaman pemantauan untuk {FIVE_DIVISIONS}, atau analis. Halaman ini menampilkan ringkasan data, grafik, status laporan, dan informasi yang membantu pengguna berwenang dalam menentukan tindak lanjut.",
686: "Kode berikut menunjukkan pemanggilan analisis untuk satu narasi laporan karena endpoint /api/ai/analyze menerima properti text.",
687: "const result = await fetch('/api/ai/analyze', { method: 'POST',",
688: "headers: { 'Content-Type': 'application/json' }, body: JSON.stringify({ text: reportText }),",
692: "Pengujian black-box pada bagian ini merupakan observasi manual dari perspektif pengguna. Skenario meliputi autentikasi, pelaporan, unggah bukti, pembaruan status, persetujuan akun, dashboard, sinkronisasi data, layanan analitik, pembatasan hak akses, dan audit. Observasi tersebut dipisahkan dari test suite repository yang dapat direproduksi.",
693: "Data uji analitik dibaca dari Google Sheets sebagai sumber operasional dengan 1.146 rekaman aktif pada snapshot evaluasi. Pengujian repository membedakan build aplikasi, unit test terfokus, lint, pengujian ML, dan pengujian Gapura RAG. Skenario black-box pada Tabel III.15 merupakan catatan observasi manual dan tidak diperlakukan sebagai bukti otomatis yang setara dengan test suite.",
697: f"Skenario pemantauan pada Tabel III.15 berlaku untuk {FIVE_DIVISIONS} dengan pola fitur yang sama. DIVISI_OCS tidak disamakan dengan kelima peran tersebut karena memakai workspace tersendiri.",
698: "Catatan pengujian manual juga mencakup skenario negatif berupa kredensial tidak terdaftar, ruas laporan wajib yang kosong, bukti melebihi batas ukuran, permintaan data di luar scope, dan layanan analitik yang tidak tersedia. Hasil observasi ini dicatat sebagai pemeriksaan eksploratif; bukti otomatis yang dapat direproduksi tetap merujuk pada build dan test suite repository.",
708: "Sistem pelaporan berhasil dirancang dengan alur transaksi yang menulis laporan ke Google Sheets terlebih dahulu dan menyinkronkan identitas serta metadata ke Supabase PostgreSQL untuk kebutuhan pembacaan aplikasi. Bukti laporan menggunakan Google Drive, sedangkan dokumen akhir DOCX/PDF menggunakan Supabase Storage melalui alur terpisah.",
709: "Alur write-first Google Sheets dan sinkronisasi metadata Supabase berpotensi mengurangi penyalinan manual dari Microsoft Excel. Namun, penghematan waktu, konsistensi sinkronisasi pada gangguan layanan, dan penurunan kesalahan belum diukur secara kuantitatif.",
711: "Empat belas skenario pada Tabel III.15 merupakan observasi manual atas alur autentikasi, pelaporan, unggah bukti, tanggapan, persetujuan akun, dashboard, sinkronisasi data, layanan analitik, pembatasan akses, audit, dan logout. Observasi tersebut memberi indikasi awal, tetapi tidak diperlakukan sebagai pengganti test suite otomatis atau bukti kinerja produksi.",
712: "Pengujian dilakukan pada lingkungan lokal dan kontainer uji sehingga kinerja beban produksi belum diketahui. Build aplikasi dan unit test terfokus berhasil, tetapi lint masih mencatat error serta warning pada snapshot audit. Pengujian ML terfokus berhasil, sedangkan discovery penuh masih mengalami error fixture dari skrip test_local.py. Gapura RAG lulus 44 dari 44 pengujian repository. Skenario black-box bersifat observasi manual dan belum mencakup UAT, pengukuran kegunaan, uji beban, atau uji keamanan. Kualitas jawaban RAG dan otorisasi dokumen per peran/divisi belum dievaluasi sebagai kontrol produksi.",
}


COMMENTS = {
155: "Koreksi arsitektur utama: implementasi menulis laporan ke Google Sheets terlebih dahulu; Supabase menerima sinkronisasi metadata/read model. Penyimpanan bukti dan dokumen final juga dipisahkan sesuai rute aktif.",
160: "English abstract synchronized with the corrected Indonesian implementation description.",
269: "Mengganti klaim Supabase write-first dengan alur transaksi yang terverifikasi pada kode.",
307: "Istilah system of record sebelumnya tidak sesuai urutan write aktual; paragraf kini membedakan sumber operasional dan read model aplikasi.",
317: "Sesi bukan stateless murni: cookie JWT didukung tabel security_sessions untuk kedaluwarsa dan revocation.",
318: "Skema live tidak memiliki ai_audit_logs; rujukan lama diganti dengan audit_logs dan ai_cache_entries yang benar-benar ada.",
422: "Model aktor diselaraskan: HC dan peran Eskalasi dikeluarkan; OS, OP, HT, UQ, dan OT tetap lima peran terpisah dengan hak yang sama; OCS terpisah.",
493: "ERD/LRS dibangun ulang dari metadata skema live Supabase dan membedakan foreign key fisik dari asosiasi logis.",
503: "Arsitektur penyimpanan dikoreksi menjadi Sheets-first, Google Drive untuk bukti, dan Supabase Storage untuk dokumen final.",
507: "Kontrak AI/RAG disesuaikan dengan rute aktif, kuota proxy, bukti sumber/halaman, dan keterbatasan ACL dokumen.",
565: "Daftar aktor implementasi diperbarui tanpa HC, tanpa peran Eskalasi, dan tanpa label payung untuk lima peran divisi.",
573: "Google Sheets adalah sumber operasional langsung layanan ML; Supabase menyimpan read model tersinkron.",
610: "Integrasi RAG dikoreksi ke gerbang sesi dan kuota OneClick serta pipeline retrieval yang benar-benar diuji.",
687: "Contoh payload diperbaiki dari data: reports menjadi text: reportText sesuai kontrak route /api/ai/analyze.",
693: "Status pengujian dipisahkan antara test suite reproducible dan observasi black-box manual.",
708: "Kesimpulan diselaraskan dengan arsitektur transaksi dan penyimpanan yang benar-benar diimplementasikan.",
712: "Keterbatasan kini mencatat hasil audit build/lint/test serta batas evaluasi RAG secara eksplisit.",
}


def set_paragraph_text(paragraph, text):
    first_rpr = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    # Remove every inline child, including stale citation/REF/TOC fields and
    # hyperlinks; keeping only pPr avoids duplicated field display text.
    for child in list(paragraph._p):
        if child.tag.endswith("}pPr"):
            continue
        paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if first_rpr is not None:
        run._r.insert(0, first_rpr)
    return run


def remove_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def replace_cell(cell, text):
    p = cell.paragraphs[0]
    run = set_paragraph_text(p, text)
    for extra in list(cell.paragraphs[1:]):
        remove_paragraph(extra)
    return run


def replace_figure(paragraph, image_path, width_inches=6.25):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))


def add_comment(doc, run, text):
    doc.add_comment(run, text=text, author=AUTHOR, initials=INITIALS)


def revise_tables(doc):
    # Tabel III.1 — non-functional requirements.
    replace_cell(doc.tables[5].rows[1].cells[2], "Sistem menggunakan login, pembatasan akses berbasis peran, validasi input, sesi database, dan audit. Kata sandi disimpan sebagai hash bcrypt; sesi normal berlaku paling lama 24 jam; unggahan bukti dibatasi 10 MB per berkas.")
    replace_cell(doc.tables[5].rows[2].cells[2], f"Halaman staf cabang, manajer cabang, {FIVE_DIVISIONS}, OCS, analis, dan super admin dirancang agar fungsi utama dapat digunakan melalui peramban desktop maupun perangkat bergerak.")
    replace_cell(doc.tables[5].rows[3].cells[2], "Laporan ditulis ke Google Sheets dan dapat dibaca kembali melalui read model Supabase, halaman laporan, serta dashboard sesuai keberhasilan sinkronisasi metadata.")
    replace_cell(doc.tables[5].rows[4].cells[2], "Sistem berkomunikasi dengan Google Sheets, Google Drive, Supabase PostgreSQL, Supabase Storage, layanan ML, dan layanan RAG melalui kontrak API.")
    replace_cell(doc.tables[5].rows[6].cells[2], "Kode dipisahkan berdasarkan halaman dan API staf, manajer, lima peran divisi, workspace OCS, analisis, administrasi, laporan, unggahan, dashboard, ML, dan RAG.")

    replace_cell(doc.tables[6].rows[4].cells[2], "Peramban modern, Google Drive untuk bukti laporan, Supabase Storage untuk dokumen akhir, layanan FastAPI untuk ML/RAG, serta cache atau hasil terakhir untuk degradasi layanan.")

    api_rows = [
        ("/api/auth/login", "Memvalidasi kredensial aplikasi dan membuat JWT cookie serta security_session.", "Pengguna terdaftar"),
        ("/api/reports/public", "Menerima pengajuan laporan publik dengan validasi dan rate limit.", "Publik"),
        ("/api/reports", "Menulis laporan ke Google Sheets lalu menyinkronkan metadata Supabase secara best-effort.", "Terautentikasi"),
        ("/api/uploads/evidence", "Mengunggah bukti ke Google Drive dan mencatat ledger bukti.", "Publik/terautentikasi sesuai mode"),
        ("/api/reports/analytics", "Menyajikan data laporan untuk dashboard dan rekap.", "Peran terautentikasi sesuai scope"),
        ("/api/ai/analyze", "Menganalisis satu narasi teks melalui orkestrasi AI.", "Semua peran terautentikasi"),
        ("/api/ai/analyze-all", "Alias GET ikhtisar analitik agregat; bukan batch POST lama.", "Terautentikasi"),
        ("/api/ai/risk/summary", "Menyajikan ringkasan risiko agregat.", "Analis/lima peran divisi"),
        ("/api/dashboards/query", "Menjalankan kueri dashboard dengan filter dan pemeriksaan akses.", "Terautentikasi sesuai scope"),
        ("/api/admin/users/approve-staff", "Menyetujui akun staf pending setelah pemeriksaan kewenangan.", "MANAGER_CABANG/SUPER_ADMIN"),
    ]
    for row, data in zip(doc.tables[7].rows[1:], api_rows):
        for cell, val in zip(row.cells[1:], data): replace_cell(cell, val)

    data_rows = [
        ("Google Sheets laporan", "Tujuan penulisan pertama dan sumber operasional layanan analitik.", "sheet_id, tanggal, maskapai, cabang, kategori, uraian, status"),
        ("users", "Akun, hash kata sandi, peran, divisi, dan status aplikasi.", "id, email, password, role, division, station_id"),
        ("ground_handling_irregularity_report", "Read model/metadata laporan yang disinkronkan dari alur Sheets.", "id, sheet_id, user_id, status, category, synced_at"),
        ("report_comments", "Komentar laporan; report_id merupakan asosiasi logis bertipe text.", "id, report_id, user_id, content, created_at"),
        ("evidence_upload_sessions & evidence_files", "Ledger sesi, kepemilikan, serta tautan bukti Google Drive.", "session_id, user_id, google_drive_file_id, report_id"),
        ("security_sessions", "Status sesi aplikasi, masa berlaku, aktivitas, dan revocation.", "id, session_id, user_id, is_revoked, expires_at"),
        ("audit_logs", "Jejak aktivitas yang benar-benar dicatat aplikasi.", "actor_id, action, entity_type, entity_id, created_at"),
        ("report_documents", "Metadata dokumen akhir DOCX/PDF pada Supabase Storage.", "report_type, report_id, docx_path, pdf_path, created_by"),
    ]
    table = doc.tables[8]
    while len(table.rows) < 9: table.add_row()
    for i, data in enumerate(data_rows, start=1):
        replace_cell(table.rows[i].cells[0], str(i))
        for cell, val in zip(table.rows[i].cells[1:], data): replace_cell(cell, val)

    tech_rows = [
        ("Tampilan aplikasi", "Next.js, React, dan TypeScript."),
        ("Data aplikasi dan sesi", "Supabase PostgreSQL dengan autentikasi aplikasi bcrypt/JWT."),
        ("Bukti laporan", "Google Drive; Supabase menyimpan ledger metadata dan pengaitan."),
        ("Laporan operasional", "Google Sheets API sebagai tujuan write-first dan sumber layanan ML."),
        ("Dokumen final", "Supabase Storage untuk berkas DOCX/PDF yang tercatat pada report_documents."),
        ("Analisis AI", "FastAPI ML service dan proxy orkestrasi OneClick."),
        ("Asisten virtual", "FastAPI Gapura RAG, Pinecone, reranking, dan jawaban bersumber."),
        ("Keamanan aplikasi", "JWT cookie, security_sessions, role checks, rate limiting, validasi, dan audit."),
    ]
    for row, data in zip(doc.tables[9].rows[1:], tech_rows):
        for cell, val in zip(row.cells[1:], data): replace_cell(cell, val)

    rag = doc.tables[18]
    replace_cell(rag.rows[12].cells[1], "Rate limit layanan per IP dapat dikonfigurasi melalui environment.")
    replace_cell(rag.rows[13].cells[1], "Proxy OneClick memerlukan sesi aplikasi dan membatasi 5 pertanyaan per akun per hari.")
    replace_cell(rag.rows[14].cells[1], "Server-Sent Events pada layanan yang mendukung streaming; respons tetap membawa metadata sumber/halaman.")

    tests = doc.tables[19]
    replace_cell(tests.rows[0].cells[3], "Status bukti")
    for row in tests.rows[1:]: replace_cell(row.cells[3], "Observasi manual")
    replace_cell(tests.rows[6].cells[1], f"{FIVE_DIVISIONS} memberi tanggapan atau status dengan pola fitur yang sama")

    maintenance = doc.tables[20]
    replace_cell(maintenance.rows[3].cells[2], "Memantau write-first ke Google Sheets, kegagalan sinkronisasi metadata Supabase, struktur kolom, dan pembacaan layanan ML.")
    replace_cell(maintenance.rows[4].cells[2], "Memastikan berkas bukti Google Drive, ledger evidence_files, dan pengaitan laporan tetap konsisten; memantau Supabase Storage untuk dokumen akhir.")
    replace_cell(maintenance.rows[5].cells[2], f"Menyesuaikan halaman staf, manajer, {FIVE_DIVISIONS}, workspace OCS, analis, dan super admin tanpa menyatukan lima peran divisi atau OCS.")
    replace_cell(maintenance.rows[6].cells[2], f"Memantau layanan AI, cache, fallback, dan kuota RAG agar fitur analitik tetap dapat digunakan oleh {FIVE_DIVISIONS} dan analis.")


def main(argv):
    if len(argv) != 4:
        raise SystemExit("usage: revise_thesis_docx.py INPUT DIAGRAM_DIR OUTPUT")
    source, diagram_dir, output = map(Path, argv[1:])
    doc = Document(source)
    paragraphs = list(doc.paragraphs)

    for idx, text in PARAGRAPH_REPLACEMENTS.items():
        run = set_paragraph_text(paragraphs[idx], text)
        if idx in COMMENTS:
            add_comment(doc, run, COMMENTS[idx])

    # Remove the obsolete DIVISI_ESKALASI requirement block.
    for idx in range(450, 455):
        remove_paragraph(paragraphs[idx])

    # Replace any remaining generic umbrella label with the five explicit actors.
    for p in list(doc.paragraphs):
        if re.search(r"divisi operasional", p.text, flags=re.I):
            new = re.sub(r"divisi operasional", FIVE_DIVISIONS, p.text, flags=re.I)
            set_paragraph_text(p, new)

    revise_tables(doc)
    # One comment on each materially revised table.
    add_comment(doc, doc.tables[5].rows[1].cells[2].paragraphs[0].runs[0], "Masa sesi disesuaikan dari tujuh hari menjadi default implementasi 24 jam; batas bukti 10 MB dipertahankan karena sesuai rute upload.")
    add_comment(doc, doc.tables[7].rows[2].cells[1].paragraphs[0].runs[0], "Kontrak API diperbarui: pengajuan publik tersedia, /analyze menerima teks tunggal, dan /analyze-all adalah alias GET overview.")
    add_comment(doc, doc.tables[8].rows[1].cells[1].paragraphs[0].runs[0], "Tabel data dikoreksi menjadi Sheets-first, Drive untuk bukti, Supabase read model, dan Storage untuk dokumen final.")
    add_comment(doc, doc.tables[19].rows[0].cells[3].paragraphs[0].runs[0], "Hasil black-box diklasifikasikan sebagai observasi manual, bukan test suite otomatis.")

    figure_map = {
        495: "01-erd-chen.png", 500: "02-lrs.png", 504: "03-architecture.png",
        510: "04-use-case.png", 515: "05-activity-login.png",
        519: "06-activity-report-submission.png", 524: "07-activity-staff-approval.png",
        529: "08-activity-ai-analysis.png", 535: "09-sequence-login.png",
        540: "10-sequence-report-submission.png", 545: "11-sequence-staff-approval.png",
        550: "12-sequence-ai-analysis.png",
    }
    for idx, filename in figure_map.items():
        replace_figure(paragraphs[idx], diagram_dir / filename)

    # Materialize captions so removing stale SEQ fields cannot renumber later figures.
    captions = {
        496:"Gambar III. 1. Entity Relationship Diagram Sistem",
        501:"Gambar III. 2. Logical Record Structure Sistem",
        505:"Gambar III. 3. Arsitektur Sistem dan Integrasi Pipeline AI",
        511:"Gambar III. 4. Use Case Diagram Keseluruhan Sistem",
        516:"Gambar III. 5. Activity Diagram: Login",
        520:"Gambar III. 6. Activity Diagram: Pengajuan Laporan",
        525:"Gambar III. 7. Activity Diagram: Persetujuan Akun Staf",
        530:"Gambar III. 8. Activity Diagram: Analisis AI dan RAG",
        536:"Gambar III. 9. Sequence Diagram: Login",
        541:"Gambar III. 10. Sequence Diagram: Pengajuan Laporan",
        546:"Gambar III. 11. Sequence Diagram: Persetujuan Akun Staf",
        551:"Gambar III. 12. Sequence Diagram: Analisis AI dan RAG",
        623:"Gambar III. 13. Tampilan Halaman Staf Cabang: Formulir Pelaporan",
        628:"Gambar III. 14. Tampilan Halaman Pemantauan: Dashboard Utama",
        633:"Gambar III. 15. Tampilan Halaman Analisis AI",
        636:"Gambar III. 16. Tampilan Halaman Asisten Virtual Berbasis RAG",
    }
    for idx, text in captions.items(): set_paragraph_text(paragraphs[idx], text)

    figure_list = [
        "Gambar III. 1. Entity Relationship Diagram Sistem 34",
        "Gambar III. 2. Logical Record Structure Sistem 35",
        "Gambar III. 3. Arsitektur Sistem dan Integrasi Pipeline AI 36",
        "Gambar III. 4. Use Case Diagram Keseluruhan Sistem 37",
        "Gambar III. 5. Activity Diagram: Login 38",
        "Gambar III. 6. Activity Diagram: Pengajuan Laporan 39",
        "Gambar III. 7. Activity Diagram: Persetujuan Akun Staf 40",
        "Gambar III. 8. Activity Diagram: Analisis AI dan RAG 41",
        "Gambar III. 9. Sequence Diagram: Login 42",
        "Gambar III. 10. Sequence Diagram: Pengajuan Laporan 43",
        "Gambar III. 11. Sequence Diagram: Persetujuan Akun Staf 44",
        "Gambar III. 12. Sequence Diagram: Analisis AI dan RAG 45",
        "Gambar III. 13. Tampilan Halaman Staf Cabang: Formulir Pelaporan 75",
        "Gambar III. 14. Tampilan Halaman Pemantauan: Dashboard Utama 76",
        "Gambar III. 15. Tampilan Halaman Analisis AI 77",
        "Gambar III. 16. Tampilan Halaman Asisten Virtual Berbasis RAG 78",
    ]
    for idx, text in zip(range(219,235), figure_list): set_paragraph_text(paragraphs[idx], text)

    table_list = [
        "Tabel III. 1. Kebutuhan Non-Fungsional Sistem 29",
        "Tabel III. 2. Kebutuhan Perangkat Keras dan Perangkat Lunak 31",
        "Tabel III. 3. Spesifikasi API Utama 45",
        "Tabel III. 4. Spesifikasi File Utama 48",
        "Tabel III. 5. Teknologi Implementasi Sistem 50",
        "Tabel III. 6. Kolom Google Sheets yang Digunakan oleh Layanan Analitik 52",
        "Tabel III. 7. Kolom Turunan (Enrichment) untuk Label Klasifikasi 57",
        "Tabel III. 8. Ringkasan Endpoint HTTP Layanan Analitik ML 60",
        "Tabel III. 9. Hasil Evaluasi Awal Model Analitik 64",
        "Tabel III. 10. Perluasan Metrik Klasifikasi dengan Majority-Class Baseline 65",
        "Tabel III. 11. Uji Sensitivitas Ambang Keyakinan (τ) pada Klasifikasi 67",
        "Tabel III. 12. Uji Sensitivitas Peringkat Risk Scoring terhadap Perubahan Bobot 68",
        "Tabel III. 13. Perbandingan Model Peramalan dan Diagnosis Residu 69",
        "Tabel III. 14. Konfigurasi Baseline Layanan RAG (Gapura RAG) 73",
        "Tabel III. 15. Skenario Pengujian Black Box 82",
        "Tabel III. 16. Rencana Pendukung dan Pemeliharaan 86",
    ]
    for idx, text in zip(range(239,255), table_list): set_paragraph_text(paragraphs[idx], text)

    # Synchronize changed high-level TOC entries after the final pagination pass.
    toc_updates = {
        192:"BAB IV PENUTUP 89", 193:"4.1 Kesimpulan 89", 194:"4.2 Saran-saran 90",
        195:"DAFTAR PUSTAKA 93", 196:"DAFTAR RIWAYAT HIDUP 97",
        197:"SURAT KETERANGAN RISET 99", 198:"BUKTI HASIL PENGECEKAN PLAGIARISME 100",
        199:"LAMPIRAN 101", 257:"A. Berita Acara Serah Terima 101",
        258:"B. Dokumentasi Berita Acara Serah Terima 104", 259:"C. Sertifikat Riset 107",
        260:"D. Press Release 109", 261:"E. Surat Persetujuan Anonimitas Identitas Mitra Riset 110",
    }
    for idx, text in toc_updates.items(): set_paragraph_text(paragraphs[idx], text)

    # Hard completion assertions for user-specified exclusions.
    combined = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    forbidden = ["DIVISI_HC", "DIVISI_ESKALASI", "DIVISI_OPERASIONAL"]
    for token in forbidden:
        if token in combined:
            raise RuntimeError(f"forbidden actor label remains: {token}")
    if re.search(r"divisi operasional", combined, flags=re.I):
        raise RuntimeError("generic umbrella label remains")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"saved {output}")


if __name__ == "__main__":
    main(sys.argv)

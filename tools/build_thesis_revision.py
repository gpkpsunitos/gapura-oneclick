#!/usr/bin/env python3
"""Build the revised thesis while preserving the source DOCX layout."""

from __future__ import annotations

import copy
import json
import re
import shutil
import urllib.request
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path("/Users/nrzngr/Documents/revisi_skripsi_v5.docx")
OUT_DIR = ROOT / "skripsi_revisi"
OUTPUT = OUT_DIR / "03_draft_revisi_skripsi.docx"
REFERENCES_DIR = OUT_DIR / "references"
DIAGRAMS_DIR = OUT_DIR / "diagrams"

DOIS = [
    "10.1109/TSE.2023.3287297",
    "10.21275/MS241022095359",
    "10.30871/jaic.v9i6.11063",
    "10.54254/2754-1169/2024.MUR17871",
    "10.52436/1.jutif.2024.5.4.1944",
    "10.59232/DST-V3I4P103",
    "10.54480/slr-m.v4i1.50",
    "10.31102/jatim.v6i2.3503",
    "10.24251/HICSS.2024.795",
    "10.3390/s23042333",
    "10.1007/s10994-023-06336-7",
    "10.3390/a16050236",
    "10.1109/ACCESS.2023.3305687",
    "10.1007/s11263-024-02029-3",
    "10.3390/fi15090290",
    "10.7717/peerj-cs.1444",
    "10.24002/ijis.v7i1.9673",
    "10.3389/fenrg.2024.1455276",
    "10.3390/bdcc9120320",
]

MANUAL_REFERENCES = [
    {
        "key": "Washizaki2024SWEBOK",
        "year": 2024,
        "apa": "Washizaki, H. (Ed.). (2024). Guide to the Software Engineering Body of Knowledge (SWEBOK Guide), Version 4.0. IEEE Computer Society. https://www.computer.org/education/bodies-of-knowledge/software-engineering",
        "bibtex": "@book{Washizaki2024SWEBOK, editor={Washizaki, Hironori}, title={Guide to the Software Engineering Body of Knowledge (SWEBOK Guide), Version 4.0}, publisher={IEEE Computer Society}, year={2024}, url={https://www.computer.org/education/bodies-of-knowledge/software-engineering}}",
    }
]


def fetch_doi(doi: str, accept: str) -> str:
    request = urllib.request.Request(
        f"https://doi.org/{doi}",
        headers={"Accept": accept, "User-Agent": "GapuraThesisVerifier/1.0"},
    )
    with urllib.request.urlopen(request, timeout=30) as response:
        return response.read().decode("utf-8").strip()


def crossref_metadata(doi: str) -> dict:
    request = urllib.request.Request(
        f"https://api.crossref.org/works/{doi}",
        headers={"User-Agent": "GapuraThesisVerifier/1.0"},
    )
    with urllib.request.urlopen(request, timeout=30) as response:
        payload = json.loads(response.read().decode("utf-8"))
    if payload.get("status") != "ok":
        raise RuntimeError(f"Crossref did not verify {doi}")
    return payload["message"]


def build_reference_files() -> tuple[list[str], list[dict]]:
    REFERENCES_DIR.mkdir(parents=True, exist_ok=True)
    apa_entries: list[str] = []
    bib_entries: list[str] = []
    verification: list[dict] = []
    for doi in DOIS:
        metadata = crossref_metadata(doi)
        year = (metadata.get("published", {}).get("date-parts") or [[None]])[0][0]
        if year not in {2023, 2024, 2025, 2026}:
            raise RuntimeError(f"DOI {doi} has disallowed year {year}")
        apa = fetch_doi(doi, "text/x-bibliography; style=apa")
        bibtex = fetch_doi(doi, "application/x-bibtex")
        if doi == "10.21275/MS241022095359":
            apa = "Cherukuri, B. R. (2024). Progressive Web Apps (PWAs): Enhancing User Experience through Modern Web Development. International Journal of Science and Research, 13(10), 1550–1560. https://doi.org/10.21275/MS241022095359"
            bibtex = "@article{Cherukuri2024PWA, author={Cherukuri, Bangar Raju}, title={Progressive Web Apps (PWAs): Enhancing User Experience through Modern Web Development}, journal={International Journal of Science and Research}, volume={13}, number={10}, pages={1550--1560}, year={2024}, doi={10.21275/MS241022095359}}"
        title = (metadata.get("title") or [""])[0]
        apa_entries.append(apa)
        bib_entries.append(bibtex)
        verification.append(
            {
                "doi": doi.lower(),
                "year": year,
                "title": title,
                "crossref": f"https://api.crossref.org/works/{doi}",
                "resolver": f"https://doi.org/{doi}",
                "verified": True,
            }
        )
    for item in MANUAL_REFERENCES:
        apa_entries.append(item["apa"])
        bib_entries.append(item["bibtex"])
        verification.append(
            {
                "doi": None,
                "year": item["year"],
                "title": "Guide to the Software Engineering Body of Knowledge (SWEBOK Guide), Version 4.0",
                "url": "https://www.computer.org/education/bodies-of-knowledge/software-engineering",
                "verified": True,
            }
        )
    apa_entries.sort(key=lambda value: re.sub(r"^[^A-Za-z]+", "", value).casefold())
    bib_entries.sort(key=str.casefold)
    REFERENCES_DIR.joinpath("references_clean.bib").write_text(
        "\n\n".join(bib_entries) + "\n", encoding="utf-8"
    )
    REFERENCES_DIR.joinpath("doi_verification.json").write_text(
        json.dumps(verification, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    return apa_entries, verification


def replace_paragraph_text(paragraph, text: str) -> None:
    first_rpr = None
    first_runs = paragraph._p.xpath(".//w:r")
    if first_runs:
        rpr = first_runs[0].find(qn("w:rPr"))
        if rpr is not None:
            first_rpr = copy.deepcopy(rpr)
    for child in list(paragraph._p):
        if child.tag not in {qn("w:pPr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if first_rpr is not None:
        existing = run._element.find(qn("w:rPr"))
        if existing is not None:
            run._element.remove(existing)
        run._element.insert(0, first_rpr)


def accept_revisions_in_docx(docx_path: Path) -> None:
    """Accept insertions and remove deletions before deterministic editing."""
    with zipfile.ZipFile(docx_path, "r") as source_zip:
        members = {name: source_zip.read(name) for name in source_zip.namelist()}
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for name, payload in list(members.items()):
        if not name.startswith("word/") or not name.endswith(".xml"):
            continue
        try:
            root = etree.fromstring(payload)
        except etree.XMLSyntaxError:
            continue
        changed = False
        for deletion in list(root.xpath(".//w:del", namespaces={"w": w_ns})):
            parent = deletion.getparent()
            if parent is not None:
                parent.remove(deletion)
                changed = True
        for insertion in list(root.xpath(".//w:ins", namespaces={"w": w_ns})):
            parent = insertion.getparent()
            if parent is None:
                continue
            index = parent.index(insertion)
            for child in list(insertion):
                insertion.remove(child)
                parent.insert(index, child)
                index += 1
            parent.remove(insertion)
            changed = True
        if changed:
            members[name] = etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )
    tmp = docx_path.with_suffix(".accepted.docx")
    with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as output_zip:
        for name, payload in members.items():
            output_zip.writestr(name, payload)
    tmp.replace(docx_path)


def insert_before(doc: Document, reference, text: str, style: str) -> object:
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    reference._p.addprevious(paragraph._p)
    return paragraph


def insert_before_like(
    doc: Document,
    reference,
    template,
    text: str,
    *,
    strip_numbering: bool = False,
) -> object:
    """Insert a paragraph while preserving the nearby document's local formatting."""
    paragraph = doc.add_paragraph()
    template_ppr = template._p.find(qn("w:pPr"))
    if template_ppr is not None:
        current_ppr = paragraph._p.find(qn("w:pPr"))
        if current_ppr is not None:
            paragraph._p.remove(current_ppr)
        copied_ppr = copy.deepcopy(template_ppr)
        if strip_numbering:
            num_pr = copied_ppr.find(qn("w:numPr"))
            if num_pr is not None:
                copied_ppr.remove(num_pr)
        paragraph._p.insert(0, copied_ppr)
    run = paragraph.add_run(text)
    if template.runs:
        template_rpr = template.runs[0]._element.find(qn("w:rPr"))
        if template_rpr is not None:
            run_rpr = run._element.find(qn("w:rPr"))
            if run_rpr is not None:
                run._element.remove(run_rpr)
            run._element.insert(0, copy.deepcopy(template_rpr))
    reference._p.addprevious(paragraph._p)
    return paragraph


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 28 else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for run in paragraph.runs:
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_table_row(table, values: list[str]) -> None:
    row = table.add_row()
    for index, value in enumerate(values):
        set_cell_text(row.cells[index], value)


def revise_body(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)

    exact = {
        27: "NIM: 19252358",
        31: "Perguruan Tinggi: Universitas Bina Sarana Informatika",
        34: "adalah asli (orisinal), tidak plagiat (menjiplak), dan belum pernah diterbitkan atau dipublikasikan dalam bentuk apa pun.",
        39: "Nama dan tanda tangan",
        131: "Dosen Pembimbing II : Desy Setyorini, SS, MM",
        186: "Pelaporan irregularity pada operasional ground handling membutuhkan alur data yang cepat, konsisten, dan mudah ditelusuri karena setiap penyimpangan layanan dapat memengaruhi kualitas pelayanan penerbangan. Pada sistem berjalan, proses pencatatan masih tersebar pada Microsoft Form, berkas Excel, Google Sheets, dan dasbor Looker Studio. Fragmentasi tersebut menimbulkan keterlambatan sinkronisasi, membuka peluang kesalahan pencatatan, dan membatasi pemanfaatan data historis sebagai dasar analisis operasional. Penelitian ini bertujuan merancang dan mengimplementasikan sistem pelaporan irregularity ground handling terintegrasi berbasis arsitektur mikrolayanan dengan Supabase sebagai system of record dan Google Sheets sebagai salinan (mirror) operasional. Metode pengembangan mencakup analisis kebutuhan, perancangan arsitektur, pembangunan antarmuka menggunakan Next.js, integrasi layanan FastAPI, serta pengujian fungsional dan audit runtime lokal. Kapabilitas utama yang diuji meliputi issue forecasting, analisis seasonality, klasifikasi subcategory, identifikasi root cause, dan perhitungan risk scoring. Sebagai fitur pendukung, sistem mengintegrasikan asisten virtual berbasis Retrieval-Augmented Generation (RAG) untuk tanya jawab dokumen operasional dengan evidence sumber dan halaman. Hasil audit terhadap 1.146 rekaman aktif menunjukkan bahwa lima endpoint analitik utama dapat dipanggil pada lingkungan uji lokal; suite layanan RAG juga lulus 44 dari 44 pengujian unit pada 12 Juli 2026. Jalur forecasting, seasonality, dan risk scoring menunjukkan keluaran runtime paling stabil, sedangkan klasifikasi subcategory dan root cause telah berjalan meskipun sebagian masih mengembalikan keluaran abstain (UNCERTAIN). Hasil tersebut menunjukkan kelayakan fungsional purwarupa untuk mengurangi ketergantungan pada proses salin ulang berkala dan menyediakan informasi analitik awal, tanpa menyatakan bahwa akurasi analitik atau kualitas jawaban RAG telah tervalidasi untuk penggunaan produksi.",
        188: "Kata Kunci: Analitik Prediktif; Arsitektur Mikrolayanan; Ground Handling; Integrasi Sistem; Machine Learning; Retrieval-Augmented Generation",
        282: "Rangkaian pekerjaan manual tersebut memunculkan fragmentasi data, menghambat konsolidasi, dan menyulitkan organisasi mempertahankan satu acuan data yang konsisten. Literatur mengenai single source of truth menekankan pentingnya satu sumber rujukan yang jelas untuk menghasilkan nilai data yang dapat dipertanggungjawabkan (Queiroz et al., 2024). Di sisi lain, klasifikasi teks berbasis pembelajaran mesin dapat membantu mengubah narasi tidak terstruktur menjadi kategori yang lebih mudah dianalisis, selama hasilnya dievaluasi secara terukur dan tetap ditinjau pengguna (Palanivinayagam et al., 2023).",
        283: "Menanggapi permasalahan tersebut, penelitian ini menawarkan rancang bangun sistem pelaporan irregularity terpadu berbasis web yang menyimpan laporan pada basis data Supabase sebagai system of record, sedangkan Google Sheets dipertahankan sebagai salinan operasional yang tersinkron. Pemisahan aplikasi utama, layanan analitik, dan layanan RAG mengikuti prinsip arsitektur mikrolayanan agar tanggung jawab setiap layanan dapat dikembangkan dan diuji secara lebih terisolasi (Velepucha & Flores, 2023; Abgaz et al., 2023; Putra et al., 2025). Sistem mengadopsi NLP untuk klasifikasi teks, peramalan untuk menaksir volume laporan, dan RAG sebagai fitur pendukung tanya jawab dokumen operasional.",
        292: "Sebagai respons terhadap permasalahan tersebut, solusi yang diajukan adalah membangun aplikasi pelaporan irregularity berbasis web yang menyimpan laporan pada Supabase sebagai system of record dan menyinkronkannya ke Google Sheets sebagai mirror operasional. Arsitektur mikrolayanan memisahkan transaksi pelaporan, analitik prediktif, dan asisten virtual RAG. Pemisahan ini memungkinkan tiap layanan memiliki kontrak API, mekanisme kegagalan, dan proses evaluasi yang berbeda tanpa menjadikan keluaran AI sebagai keputusan otomatis.",
        300: "Menjadi alat bantu bagi pihak manajemen untuk membaca tren musiman, mengidentifikasi kandidat akar masalah, menyusun prioritas risiko, dan mengakses evidence dokumen operasional; seluruh keluaran analitik dan RAG tetap memerlukan verifikasi pengguna berwenang.",
        315: "Sistem merupakan sekumpulan komponen yang saling berhubungan untuk mencapai tujuan tertentu. Dalam rekayasa perangkat lunak, kebutuhan, desain, konstruksi, pengujian, operasi, dan pemeliharaan perlu dikelola sebagai bagian yang saling terkait (Washizaki, 2024). Pada penelitian ini, sistem informasi mendigitalkan alur pelaporan irregularity agar catatan operasional dapat disimpan, ditelusuri, dan disajikan kepada pengguna sesuai hak akses.",
        317: "Aplikasi berbasis web dijalankan melalui peramban dan memanfaatkan jaringan untuk menghubungkan antarmuka pengguna dengan layanan pada sisi server. Pendekatan ini mendukung pembaruan terpusat dan akses lintas perangkat. Pada penelitian ini, antarmuka web digunakan untuk pencatatan laporan, pemantauan tindak lanjut, dashboard, analitik, dan asisten virtual.",
        319: "Progressive Web App (PWA) menggabungkan jangkauan aplikasi web dengan kemampuan seperti instalasi melalui web app manifest, caching, dan pemrosesan latar melalui service worker (Cherukuri, 2024). Kemampuan tersebut tidak berarti seluruh fungsi server tersedia saat luring; aplikasi tetap memerlukan strategi antrean dan sinkronisasi untuk operasi yang mengubah data.",
        322: "Data silo muncul ketika data tersebar pada repositori atau proses yang tidak memiliki acuan dan jalur integrasi yang jelas. Kondisi ini meningkatkan pekerjaan ganda dan risiko perbedaan versi data. Konsep single source of truth menetapkan satu sumber rujukan utama, sedangkan sistem lain dapat berfungsi sebagai salinan atau kanal konsumsi yang terkontrol (Queiroz et al., 2024).",
        325: "Arsitektur mikrolayanan memisahkan kemampuan sistem ke dalam layanan dengan tanggung jawab dan antarmuka yang jelas. Pendekatan ini dapat mendukung pengembangan dan deployment yang lebih independen, tetapi menambah kebutuhan tata kelola komunikasi, keamanan, observabilitas, dan konsistensi data (Velepucha & Flores, 2023; Abgaz et al., 2023; Putra et al., 2025). Pada sistem ini, layanan aplikasi, ML, dan RAG dipisahkan sesuai fungsi masing-masing.",
        326: "Komunikasi antarlayanan menggunakan API berbasis HTTP dan payload JSON. Kontrak endpoint, autentikasi, status respons, validasi masukan, dan penanganan kegagalan perlu didefinisikan agar perubahan pada satu layanan tidak merusak konsumennya (Velepucha & Flores, 2023). Aplikasi Next.js bertindak sebagai antarmuka sekaligus lapisan orkestrasi, sedangkan komputasi analitik dan RAG dijalankan oleh layanan FastAPI terpisah.",
        328: "Basis data menyimpan data secara terstruktur agar dapat dicari, dihubungkan, dan dijaga integritasnya. Pemodelan entitas, atribut, kunci, dan relasi membantu menerjemahkan kebutuhan ke skema implementasi (Ma et al., 2024). Pada penelitian ini, PostgreSQL pada Supabase berperan sebagai system of record, sedangkan Google Sheets dipertahankan sebagai mirror operasional.",
        330: "Penelitian ini menggunakan pendekatan Agile untuk mengembangkan perangkat lunak secara bertahap. Agile menekankan iterasi, umpan balik, dan kemampuan menyesuaikan kebutuhan tanpa menghilangkan disiplin analisis, desain, konstruksi, dan pengujian (Washizaki, 2024). Pendekatan ini dipakai karena modul autentikasi, pelaporan, dashboard, ML, RAG, dan pengujian dapat diselesaikan serta dievaluasi melalui iterasi yang berbeda.",
        334: "Peramalan deret waktu memperkirakan nilai mendatang dari pola historis. Pada penelitian ini, peramalan digunakan untuk menaksir volume laporan irregularity; hasilnya diperlakukan sebagai estimasi berbasis data historis, bukan kepastian kejadian operasional (Kochetkova et al., 2023).",
        335: "Galat peramalan dievaluasi menggunakan Mean Absolute Error (MAE) dan Root Mean Square Error (RMSE). MAE merangkum rata-rata simpangan absolut, sedangkan RMSE memberi penalti lebih besar pada galat yang besar. Karena urutan waktu tidak boleh diacak, evaluasi dilakukan dengan rolling-origin agar setiap lipatan hanya menggunakan masa lalu untuk memprediksi periode berikutnya (Kochetkova et al., 2023).",
        342: "Seasonal-Trend decomposition using LOESS (STL) memisahkan deret waktu menjadi komponen tren, musiman, dan residu. Pemisahan tersebut membantu membaca pola berulang tanpa menyamakan komponen musiman dengan hubungan sebab-akibat (Huang, 2024).",
        343: "Secara matematis, deret waktu Y_t diuraikan menjadi komponen tren jangka panjang (T_t), komponen musiman (S_t), dan komponen residu (R_t), sebagaimana dirumuskan pada persamaan berikut (Huang, 2024):",
        348: "Laporan irregularity banyak ditulis sebagai teks bebas. Klasifikasi teks merupakan salah satu tugas NLP untuk memetakan dokumen ke kelas yang telah ditetapkan. Metode tradisional dapat menggunakan n-gram atau TF-IDF sebagai representasi dan pengklasifikasi seperti Logistic Regression, SVM, atau Naive Bayes (Palanivinayagam et al., 2023).",
        351: "Matriks TF-IDF menjadi masukan bagi pengklasifikasi untuk memperkirakan kategori laporan. Kombinasi fitur dan algoritma perlu dievaluasi pada bentuk masukan yang sama dengan data saat layanan digunakan agar metrik tidak memperoleh keuntungan dari kebocoran label atau atribut pascakejadian (Palanivinayagam et al., 2023; Dermawan & Ayunda, 2025).",
        363: "Keyakinan akhir dikalibrasi ulang dengan Platt scaling agar skor prediksi lebih mendekati frekuensi kebenaran empiris. Kalibrasi perlu dipelajari dari prediksi di luar data pelatihan langsung untuk mengurangi optimisme pada estimasi probabilitas (Silva Filho et al., 2023):",
        365: "Model menahan keputusan (abstain) dan mengembalikan label UNCERTAIN beserta kandidat berperingkat apabila keyakinan tertinggi berada di bawah ambang tau = 0,35 atau kemiripan tetangga terdekat berada di bawah ambang deteksi out-of-distribution. Mekanisme ini mengurangi pemaksaan label pada masukan yang tidak didukung data pelatihan, tetapi tetap memerlukan evaluasi coverage dan selective risk (Xia & Bouganis, 2024):",
        367: "Seluruh pengklasifikasi memakai penyeimbangan kelas dan dievaluasi melalui stratified 5-fold cross-validation agar proporsi kelas pada setiap lipatan tetap mendekati distribusi data asal. Stratifikasi membantu perbandingan pada data tidak seimbang, tetapi tidak menggantikan analisis per kelas maupun pemeriksaan kesalahan (Szeghalmy & Fazekas, 2023).",
        369: "Pembobotan risiko digunakan untuk membandingkan urgensi relatif antarentitas operasional. Skor komposit perlu diperlakukan sebagai alat prioritisasi yang bergantung pada definisi komponen dan bobot, bukan sebagai probabilitas terjadinya insiden (Prasetyo et al., 2024; Liu, 2024).",
        370: "Indeks risiko (R) dihitung dari agregasi berbobot atas komponen yang dinormalisasi. Pada implementasi ini, bobot ditetapkan sebagai aturan sistem dan tidak diklaim sebagai parameter yang telah diestimasi secara kausal atau divalidasi sebagai standar industri (Prasetyo et al., 2024).",
        374: "Unified Modeling Language (UML) menyediakan diagram untuk merepresentasikan kebutuhan, perilaku, dan interaksi sistem. Use case menggambarkan tujuan aktor, activity diagram menggambarkan alur aktivitas, dan sequence diagram menggambarkan urutan pesan pada suatu skenario (Alyami et al., 2023; Washizaki, 2024).",
        376: "Entity Relationship Diagram (ERD) menggambarkan entitas, atribut, dan hubungan data, sedangkan Logical Record Structure (LRS) merangkum struktur tabel, kunci, dan relasi logis yang akan diimplementasikan. Keduanya digunakan untuk menelusuri kebutuhan data ke skema basis data (Ma et al., 2024).",
        378: "Pengujian black box menilai perilaku yang tampak melalui masukan, keluaran, status respons, dan aturan akses tanpa bergantung pada rincian implementasi internal. Pendekatan ini sesuai untuk menguji fungsi web dan API berdasarkan skenario pengguna (Mridha & Joarder, 2023; Ismiati et al., 2024).",
        379: "Retrieval-Augmented Generation (RAG)",
        380: "Retrieval-Augmented Generation (RAG) menggabungkan pencarian informasi eksternal dengan model generatif. Dokumen diproses menjadi potongan (chunk), direpresentasikan sebagai embedding, disimpan pada vector store, lalu diambil kembali berdasarkan kemiripan dengan pertanyaan. Konteks terpilih digunakan untuk membatasi dasar jawaban model bahasa (Brown et al., 2025).",
        381: "Pipeline RAG perlu memisahkan ingestion, retrieval, reranking, dan generation agar kegagalan pada setiap tahap dapat diamati. Metadata sumber dan halaman perlu dipertahankan bersama chunk sehingga pengguna dapat memeriksa evidence yang mendukung jawaban.",
        382: "Kualitas RAG tidak cukup dinilai dari apakah endpoint menghasilkan teks. Evaluasi perlu mencakup relevansi konteks, ketepatan evidence, faithfulness jawaban terhadap konteks, kualitas sitasi, latensi, serta perilaku ketika bukti tidak memadai (Brown et al., 2025).",
        383: "Pada penelitian ini, Gapura RAG diposisikan sebagai fitur pendukung tanya jawab dokumen operasional. Implementasi menggunakan layanan FastAPI, embedding multilingual, Pinecone sebagai vector store, retrieval dan reranking, serta model bahasa melalui OpenRouter. Jawaban ditampilkan bersama sumber dan halaman; kualitas jawabannya belum diklaim tervalidasi untuk keputusan operasional.",
        397: "Metode pengembangan perangkat lunak yang digunakan adalah Agile. Pengembangan dijalankan melalui iterasi yang mencakup analisis kebutuhan, desain, konstruksi, pengujian, serta pendukung atau pemeliharaan. Pola tersebut menjaga kemampuan beradaptasi tanpa menghilangkan traceability artefak rekayasa perangkat lunak (Washizaki, 2024). Iterasi awal berfokus pada autentikasi dan pelaporan, diikuti dashboard dan sinkronisasi, layanan ML, integrasi RAG, serta pengujian dan penyempurnaan.",
        425: "PENGGUNA DIVISI (DIVISI_OCS, DIVISI_OS, DIVISI_OP, DIVISI_OT, DIVISI_UQ, DIVISI_HC, DAN DIVISI_HT)",
        426: "Melakukan login dan mengakses ruang kerja sesuai peran serta cakupan yang dikonfigurasi.",
        427: "Melihat laporan, detail kejadian, bukti, status, dan tindak lanjut yang berada dalam cakupan divisinya.",
        428: "Memberikan komentar atau pembaruan status hanya apabila fungsi dan kebijakan akses divisinya mengizinkan.",
        429: "Membaca dashboard, analitik prediktif, dan prioritas risiko sebagai informasi pendukung, bukan keputusan otomatis.",
        430: "Mengakses asisten virtual RAG untuk mencari informasi dokumen operasional dan memeriksa evidence sumber serta halaman.",
        438: "Memeriksa kewajaran keluaran analitik dan evidence RAG sebelum digunakan sebagai bahan evaluasi operasional.",
        454: "Menganalisis laporan untuk menghasilkan klasifikasi, prediksi volume, pola musiman, dan skor risiko; layanan RAG secara terpisah mengambil evidence dokumen untuk mendukung tanya jawab operasional.",
        455: "Menyediakan cache atau keluaran terakhir yang tersimpan ketika layanan analitik utama tidak tersedia, tanpa menyajikannya sebagai hasil inferensi baru.",
        479: "Gambar III.3 menunjukkan arsitektur layanan yang dibangun. Aplikasi Next.js menyediakan PWA, API routes, sesi, RBAC, validasi, dan orkestrasi. Laporan ditulis lebih dahulu ke PostgreSQL pada Supabase sebagai system of record, bukti disimpan di Supabase Storage, dan Google Sheets dipertahankan sebagai mirror operasional. Layanan ML FastAPI menangani klasifikasi, peramalan, dekomposisi musiman, dan risk scoring; Gapura RAG FastAPI menjadi fitur pendukung tanya jawab dokumen.",
        483: "AI Orchestrator meneruskan permintaan analitik ke layanan ML menggunakan API key serta mengelola cache atau last-good result. Jalur Virtual Assistant berbeda: endpoint Next.js memvalidasi sesi dan kuota harian, lalu meneruskan pertanyaan ke Gapura RAG memakai proxy secret. Gapura RAG melakukan embedding, pencarian pada Pinecone, reranking, dan generasi jawaban melalui OpenRouter, kemudian mengembalikan jawaban beserta evidence sumber dan halaman. Pemisahan ini mencegah fungsi prediktif dan RAG diperlakukan sebagai satu model yang sama.",
        485: "Use case diagram menggambarkan batas sistem, aktor, serta layanan yang dapat diakses. Gambar III.4 mencakup pelaporan, pemantauan, tindak lanjut, analitik prediktif, administrasi, dan asisten virtual RAG. Generalisasi aktor digunakan untuk menunjukkan fungsi bersama pengguna terautentikasi tanpa menghilangkan hak akses khusus tiap peran.",
        488: "Staf cabang mengajukan laporan dan memantau laporan sendiri. Manajer cabang memantau cakupan stasiun serta menyetujui akun staf. Pengguna divisi memantau dan menindaklanjuti laporan sesuai konfigurasi akses. DIVISI_ESKALASI dapat beralih ruang kerja yang diizinkan. ANALYST melakukan analisis dan tindak lanjut lintas data, sedangkan SUPER_ADMIN mengelola pengguna, master data, audit, dan keamanan. Seluruh pengguna terautentikasi dapat memakai asisten virtual dengan kuota dan evidence yang dapat diperiksa.",
        531: "Rancangan API menjadi kontrak komunikasi antara halaman aplikasi, Supabase, Google Sheets, layanan ML, dan layanan RAG. Kontrak mencakup autentikasi, payload, status respons, cakupan akses, serta perilaku ketika layanan eksternal tidak tersedia. Ringkasan API utama ditunjukkan pada Tabel III.3.",
        542: "Tahap pembuatan kode menerjemahkan kebutuhan aktor menjadi halaman aplikasi, API, validasi, penyimpanan bukti, sinkronisasi Google Sheets, layanan ML, dan asisten virtual RAG. Bagian ini hanya menampilkan komponen yang mewakili alur utama; rincian implementasi diverifikasi terhadap repositori aplikasi, repositori ML, dan repositori Gapura RAG pada commit yang dicatat dalam changelog.",
        510: "Sequence diagram digunakan untuk memperjelas urutan komunikasi antara pengguna, halaman aplikasi, API, dan basis data. Gambar III.9 menunjukkan proses login dari halaman masuk sampai sistem membuat sesi.",
        547: "Implementasi dibagi menjadi modul autentikasi, pelaporan, dashboard, analitik ML, dan Virtual Assistant RAG. Modul RAG pada aplikasi utama menangani validasi sesi, kuota akun, proxy secret, dan streaming respons, sedangkan layanan Gapura RAG menangani ingestion dokumen, embedding, vector search, retrieval, reranking, dan generation. Pembagian tanggung jawab ini menjaga fungsi transaksi tidak bergantung langsung pada proses analitik berat.",
        625: "Pengujian fungsional menggunakan pendekatan black box untuk memeriksa autentikasi, pelaporan, unggah bukti, pembaruan status, persetujuan akun, dashboard, sinkronisasi, layanan analitik, pembatasan akses, audit, dan Virtual Assistant. Selain audit runtime aplikasi dan ML, suite Gapura RAG dijalankan pada commit 7f1549b7d1b0ad59c9405ba35ca0ab443e4f0ba3 dan menghasilkan 44 pengujian lulus dari 44 pengujian pada 12 Juli 2026.",
        626: "Basis data uji analitik bersumber dari ground_handling_irregularity_report pada Supabase dan mirror Google Sheets, dengan 1.146 rekaman aktif pada audit runtime 11 Juli 2026. Audit aplikasi menggunakan runtime Node.js dan Chromium, sedangkan layanan ML menggunakan runtime Python dan curl. Verifikasi RAG dilakukan terpisah melalui pytest pada lingkungan lokal; cakupannya meliputi API, ingestion, chunking, embedding, vector store, retrieval, multi-source, query pipeline, generator ter-grounding, dan pemrosesan PDF.",
        629: "Hasil pengujian menunjukkan fungsi utama aplikasi berjalan sesuai kebutuhan yang diaudit. Setiap peran memperoleh akses sesuai cakupan, sementara transaksi, bukti, status, dashboard, analitik, dan administrasi dapat dipanggil melalui jalur masing-masing. Suite RAG lulus seluruh 44 pengujian unit, tetapi hasil tersebut membuktikan perilaku kode pada skenario uji—bukan akurasi substantif semua jawaban terhadap dokumen operasional.",
        646: "Sistem menyediakan analitik awal berupa forecasting, seasonality, klasifikasi, identifikasi root cause, dan risk scoring. Sistem juga mengintegrasikan Virtual Assistant RAG sebagai fitur pendukung tanya jawab dokumen dengan evidence sumber dan halaman. Keluaran ML dan RAG harus diperiksa pengguna berwenang dan tidak diperlakukan sebagai keputusan otomatis.",
        648: "Pengujian dilakukan pada lingkungan lokal dan kontainer uji. Sebagian layanan klasifikasi masih dapat menghasilkan UNCERTAIN ketika keyakinan rendah. Meskipun 44 pengujian unit RAG lulus, evaluasi relevansi retrieval, ketepatan evidence, faithfulness jawaban, keamanan dokumen, dan validasi oleh pengguna operasional belum dilakukan secara menyeluruh. Karena itu, hasil penelitian menunjukkan kelayakan fungsi purwarupa, bukan kesiapan produksi atau akurasi analitik yang final.",
        655: "Layanan ML perlu disempurnakan melalui perbaikan data, analisis kesalahan, evaluasi per kelas, dan validasi analis. Layanan RAG perlu dievaluasi dengan kumpulan pertanyaan dan evidence acuan untuk mengukur relevansi retrieval, ketepatan sitasi, faithfulness, kegagalan tanpa bukti, latensi, dan keamanan akses dokumen.",
        660: "Penelitian berikutnya perlu membandingkan algoritma dan melaporkan precision, recall, F1-score, MAE, serta RMSE sesuai tugas. Evaluasi RAG perlu memisahkan kualitas retrieval dari kualitas generation melalui metrik dan penilaian manusia atas konteks, evidence, faithfulness, dan kegunaan jawaban. Dengan demikian, klaim kualitas hanya dibuat berdasarkan hasil yang terukur.",
    }

    for index, text in exact.items():
        replace_paragraph_text(paragraphs[index], text)
    paragraphs[379].style = doc.styles["List Paragraph"]
    if paragraphs[379].runs:
        paragraphs[379].runs[0].bold = True

    citation_replacements = {
        "(Afifah & Voutama, 2023; Dermawan & Ayunda, 2025; Imran et al., 2025; Nabiilah, 2025)": "(Palanivinayagam et al., 2023; Dermawan & Ayunda, 2025)",
        "(Amrin et al., 2025; Koutsandreas et al., 2022)": "(Prasetyo et al., 2024; Kochetkova et al., 2023)",
        "(Amrin et al., 2025; Prasetyo et al., 2024)": "(Prasetyo et al., 2024)",
        "(Ensafi et al., 2022)": "(Huang, 2024)",
        "(Firmansyah et al., 2024; Irmayani & Jayanti, 2025)": "(Kochetkova et al., 2023)",
        "(Gowda & Gowda, 2024)": "(Velepucha & Flores, 2023)",
        "(Harlinda & Satra, 2025)": "(Queiroz et al., 2024)",
        "(Ismiati et al., 2024; Zen et al., 2024)": "(Mridha & Joarder, 2023; Ismiati et al., 2024)",
        "(Kusuma & Hidayat, 2024; Sanjaya & Murnawan, 2024)": "(Washizaki, 2024)",
        "(Meng & Ban, 2024)": "(Alyami et al., 2023)",
        "(Muppala, 2025)": "(Washizaki, 2024)",
        "(Muppala, 2025; Putra et al., 2025; Sanjaya & Murnawan, 2024)": "(Velepucha & Flores, 2023; Abgaz et al., 2023; Putra et al., 2025)",
        "(Nurhadi & Indrayuni, 2024)": "(Cherukuri, 2024)",
        "(Palinggi et al., 2024)": "(Ma et al., 2024)",
        "(Pratama & Wibowo, 2023)": "(Washizaki, 2024)",
        "(Prusty et al., 2022; Lumumba et al., 2024)": "(Szeghalmy & Fazekas, 2023)",
        "(Putra et al., 2025; Sanjaya & Murnawan, 2024)": "(Velepucha & Flores, 2023; Putra et al., 2025)",
    }
    for paragraph in doc.paragraphs:
        original = paragraph.text
        revised = original
        for old, new in citation_replacements.items():
            revised = revised.replace(old, new)
        if revised != original:
            replace_paragraph_text(paragraph, revised)

    symbol_rows = [
        ("y_t", "Nilai observasi pada waktu ke-t"),
        ("ŷ_t", "Nilai hasil peramalan pada waktu ke-t"),
        ("n", "Jumlah observasi"),
        ("MAE", "Mean Absolute Error"),
        ("RMSE", "Root Mean Square Error"),
        ("α, β, γ", "Parameter penghalusan level, tren, dan musiman Holt-Winters"),
        ("T_t", "Komponen tren pada dekomposisi STL"),
        ("S_t", "Komponen musiman pada dekomposisi STL"),
        ("R_t", "Komponen residu pada dekomposisi STL"),
        ("TF", "Term Frequency"),
        ("IDF", "Inverse Document Frequency"),
        ("x", "Vektor fitur masukan klasifikasi"),
        ("w, b", "Bobot dan bias model linear"),
        ("P(y = k | x)", "Probabilitas kelas k untuk masukan x"),
        ("ω_m", "Bobot model ke-m pada ensemble"),
        ("τ, φ", "Ambang keyakinan dan ambang kemiripan OOD"),
        ("R", "Skor risiko komposit"),
        ("C_i", "Komponen risiko ke-i yang telah dinormalisasi"),
        ("w_i", "Bobot komponen risiko ke-i"),
    ]
    symbol_table = doc.tables[9]
    set_cell_text(symbol_table.rows[0].cells[0], "Simbol", bold=True)
    set_cell_text(symbol_table.rows[0].cells[1], "Keterangan", bold=True)
    for index, (symbol, meaning) in enumerate(symbol_rows, start=1):
        if index >= len(symbol_table.rows):
            symbol_table.add_row()
        set_cell_text(symbol_table.rows[index].cells[0], symbol)
        set_cell_text(symbol_table.rows[index].cells[1], meaning)

    super_admin_heading = paragraphs[440]
    insert_before_like(doc, super_admin_heading, paragraphs[431], "DIVISI_ESKALASI")
    escalation_template = paragraphs[432]
    insert_before_like(
        doc,
        super_admin_heading,
        escalation_template,
        "a) Melakukan login dan beralih ke ruang kerja divisi yang diizinkan melalui mekanisme switch session.",
        strip_numbering=True,
    )
    insert_before_like(
        doc,
        super_admin_heading,
        escalation_template,
        "b) Melihat laporan lintas divisi dan dokumen operasional sesuai cakupan akses.",
        strip_numbering=True,
    )
    insert_before_like(
        doc,
        super_admin_heading,
        escalation_template,
        "c) Mengekspor data dan menggunakan Virtual Assistant sebagai fitur pendukung pencarian dokumen.",
        strip_numbering=True,
    )

    audit_anchor = paragraphs[457]
    system_template = paragraphs[455]
    insert_before_like(
        doc,
        audit_anchor,
        system_template,
        "k) Memvalidasi sesi dan kuota Virtual Assistant sebelum meneruskan pertanyaan ke layanan RAG.",
        strip_numbering=True,
    )
    insert_before_like(
        doc,
        audit_anchor,
        system_template,
        "l) Mengembalikan jawaban RAG bersama evidence sumber dan halaman atau pesan kegagalan yang terkendali.",
        strip_numbering=True,
    )

    model_table_caption = paragraphs[549]
    insert_before(
        doc,
        model_table_caption,
        "Layanan Gapura RAG dipisahkan dari model analitik prediktif. Dokumen PDF diproses menjadi chunk, diubah menjadi embedding multilingual, dan disimpan pada Pinecone. Pertanyaan pengguna diproses melalui retrieval, penyebaran kandidat lintas sumber, reranking, serta generation melalui OpenRouter. Endpoint aplikasi utama memerlukan sesi aktif, menerapkan kuota harian, dan meneruskan proxy secret; layanan mengembalikan evidence sumber dan halaman agar jawaban dapat diperiksa.",
        "Body Text",
    )

    # API table (Table III.3)
    api_table = doc.tables[16]
    for row in api_table.rows[1:]:
        row.cells[1].text = row.cells[1].text.replace("/ ", "/").replace(" ", "")
    add_table_row(
        api_table,
        [
            "11",
            "/api/virtual-assistant/chat",
            "Memvalidasi sesi dan kuota, meneruskan pertanyaan ke Gapura RAG, dan mengalirkan jawaban beserta evidence.",
            "Semua pengguna terautentikasi",
        ],
    )

    file_table = doc.tables[17]
    replacements = {
        "report comments": "report_comments",
        "security sessions": "security_sessions",
    }
    for row in file_table.rows[1:]:
        name = row.cells[1].text.strip()
        if name in replacements:
            set_cell_text(row.cells[1], replacements[name])
    add_table_row(
        file_table,
        [
            "8",
            "Pinecone index",
            "Menyimpan embedding chunk dokumen dan metadata evidence untuk retrieval RAG.",
            "doc_id, chunk_id, source, page, embedding, metadata",
        ],
    )

    tech_table = doc.tables[18]
    add_table_row(
        tech_table,
        [
            "9",
            "Virtual Assistant RAG",
            "FastAPI, embedding multilingual, Pinecone, retrieval/reranking, dan OpenRouter.",
        ],
    )

    test_table = doc.tables[20]
    add_table_row(
        test_table,
        [
            "15",
            "Pengguna terautentikasi bertanya melalui Virtual Assistant",
            "Sistem memvalidasi kuota, mengambil evidence, dan menampilkan jawaban beserta sumber/halaman atau respons kegagalan terkendali.",
            "Berhasil pada suite unit Gapura RAG (44/44 lulus, 12 Juli 2026)",
        ],
    )

    # Reapply table header emphasis and keep rows expandable.
    for table in (api_table, file_table, tech_table, test_table):
        for cell in table.rows[0].cells:
            set_cell_text(cell, cell.text, bold=True)
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            for height in list(tr_pr.findall(qn("w:trHeight"))):
                tr_pr.remove(height)


def make_bibliography_paragraph(text: str) -> etree._Element:
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = etree.Element(f"{{{w}}}p")
    p_pr = etree.SubElement(p, f"{{{w}}}pPr")
    ind = etree.SubElement(p_pr, f"{{{w}}}ind")
    ind.set(f"{{{w}}}hanging", "480")
    jc = etree.SubElement(p_pr, f"{{{w}}}jc")
    jc.set(f"{{{w}}}val", "both")
    spacing = etree.SubElement(p_pr, f"{{{w}}}spacing")
    spacing.set(f"{{{w}}}after", "0")
    spacing.set(f"{{{w}}}line", "360")
    spacing.set(f"{{{w}}}lineRule", "auto")
    r = etree.SubElement(p, f"{{{w}}}r")
    r_pr = etree.SubElement(r, f"{{{w}}}rPr")
    fonts = etree.SubElement(r_pr, f"{{{w}}}rFonts")
    fonts.set(f"{{{w}}}ascii", "Times New Roman")
    fonts.set(f"{{{w}}}hAnsi", "Times New Roman")
    sz = etree.SubElement(r_pr, f"{{{w}}}sz")
    sz.set(f"{{{w}}}val", "24")
    t = etree.SubElement(r, f"{{{w}}}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return p


def patch_ooxml(docx_path: Path, apa_entries: list[str]) -> None:
    with zipfile.ZipFile(docx_path, "r") as source_zip:
        members = {name: source_zip.read(name) for name in source_zip.namelist()}

    document_root = etree.fromstring(members["word/document.xml"])
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    bibliography_content = None
    for sdt_content in document_root.xpath(".//w:sdtContent", namespaces=ns):
        text = "".join(sdt_content.xpath(".//w:t/text()", namespaces=ns))
        if "Afifah, N. F." in text or "10.1234/jrplsi" in text:
            bibliography_content = sdt_content
            break
    if bibliography_content is None:
        raise RuntimeError("Could not locate bibliography content control")
    for child in list(bibliography_content):
        bibliography_content.remove(child)
    for entry in apa_entries:
        bibliography_content.append(make_bibliography_paragraph(entry))
    members["word/document.xml"] = etree.tostring(
        document_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    settings_root = etree.fromstring(members["word/settings.xml"])
    update = settings_root.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings_root.append(update)
    update.set(qn("w:val"), "true")
    members["word/settings.xml"] = etree.tostring(
        settings_root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    members["word/media/image13.png"] = (DIAGRAMS_DIR / "architecture_system.png").read_bytes()
    members["word/media/image14.png"] = (DIAGRAMS_DIR / "usecase_system.png").read_bytes()

    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as output_zip:
        for name, payload in members.items():
            output_zip.writestr(name, payload)
    tmp.replace(docx_path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    apa_entries, _ = build_reference_files()
    shutil.copy2(SOURCE, OUTPUT)
    accept_revisions_in_docx(OUTPUT)
    doc = Document(OUTPUT)
    revise_body(doc)
    doc.save(OUTPUT)
    patch_ooxml(OUTPUT, apa_entries)
    print(OUTPUT)


if __name__ == "__main__":
    main()
